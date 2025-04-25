# app/output/docx_writer.py

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def create_docx_from_text(text, output_path, language='en'):
    """
    Create a .docx document from extracted text.

    Args:
        text (str): The text to write into the document.
        output_path (str): Path where the docx file will be saved.
        language (str): 'en' for English, 'ar' for Arabic (affects alignment and font).

    Returns:
        str: Path to the created document.
    """
    if not text.strip():
        raise ValueError("Cannot create a document with empty text.")

    document = Document()

    # Set default font
    if language == 'ar':
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            p = document.add_paragraph()
            run = p.add_run(para.strip())
            run.font.size = Pt(12)

            if language == 'ar':
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    document.save(output_path)
    return output_path

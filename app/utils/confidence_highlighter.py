from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def get_highlight_color(confidence):
    """
    Determines the highlight color based on confidence score.
    """
    if confidence <= 50:
        return "FF0000"  # Red
    elif confidence <= 75:
        return "FFA500"  # Orange
    elif confidence <= 90:
        return "FFFF00"  # Yellow
    else:
        return "00FF00"  # Green

def add_colored_text(paragraph, text, confidence):
    """
    Adds a run to the paragraph with highlight color based on confidence.
    """
    run = paragraph.add_run(text)

    # Apply color highlight
    highlight_color = get_highlight_color(confidence)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), highlight_color)
    run._r.get_or_add_rPr().append(shading_elm)

    # Set font properties
    font = run.font
    font.size = None  # Keep default size
    font.name = 'Arial'

def create_highlighted_document(results, output_path, rtl=False):
    """
    Creates a Word document with highlighted text based on confidence scores.
    Args:
        results (list): List of dictionaries containing 'text' and 'confidence'.
        output_path (str): Path to save the Word document.
        rtl (bool): If True, sets the paragraph direction to RTL (for Arabic).
    """
    document = Document()
    paragraph = document.add_paragraph()

    if rtl:
        # Set RTL direction for Arabic paragraphs
        paragraph.paragraph_format.alignment = 3  # Right alignment
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    for item in results:
        text = item['text']
        confidence = item['confidence']

        if text.strip() != "":
            add_colored_text(paragraph, text + ' ', confidence)

    document.save(output_path)
